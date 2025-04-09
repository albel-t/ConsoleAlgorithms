from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE

import glob
import os

def addCode(name, paragraph):
    p = paragraph.add_run(name)
    p.font.name = 'Consolas'
    p.font.size = Pt(12)
    return p

def addTask(task, paragraph):
    p = paragraph.add_run(task)
    p.font.name = 'Times_New_Roman'
    p.font.size = Pt(14)
    return p

def addCodeTitle(code, paragraph):
    p = addTask(code, paragraph)
    p.font.bold = True
    return p

def addTaskTitle(text, paragraph):
    p = addTask(text, paragraph)
    p.font.bold = True
    return p


#visual studio
# path_to_project = input()
path_to_project = "C:/Users/Bill/source/repos/ConsoleAlgorithms"
path_to_template = "C:/Users/Bill/source/repos/ConsoleAlgorithms/algorithms_template.docx"
name_solution = "ConsoleAlgorithms"
name_laba = "laba_1"
# �������� ���������
doc = Document(path_to_template)


# ������ ����� ������ �� ���������
style = doc.styles['Normal']
style.font.name = 'Times_New_Roman'
style.font.size = Pt(14)
p = doc.add_paragraph(" ")

addTaskTitle("���� ������:\n", p)
print("������� ���� ������:")
addTask(input() + "\n", p)

addTaskTitle("������e:\n", p)
print("������� ������e:")
addTask(input() + "\n", p)

laba_file = path_to_project + '/' + name_solution + '/' + name_solution + '.cpp'

file_text = open(laba_file, "r", encoding="utf8").read()

addTaskTitle("�������:\n", p)
# print(file_text, file_text.find("/***") , file_text.find("***/"), file_text[file_text.find("/***") : file_text.find("***/")])
addCode(file_text[file_text.find("/****")+5 : file_text.find("****/")], p) #����� �����������
print("������� - ", laba_file)

addCodeTitle("Code: " + name_solution + '.cpp\n', p)
addCode(file_text[file_text.find("****/")+5 : ], p) #����� �����������
print("write: ", name_solution + '.cpp')


laba_file = path_to_project + '/' + name_solution + '/' + name_laba
for filename in glob.glob(laba_file + ".*"):    #�������� ������������ ����������� ����� ��������� ".../v[������].txt"
    addCodeTitle("\nCode: " + filename[filename.find(name_laba) : ] + '\n', p)
    addCode(open(filename, "r").read(), p) #����� �����������
    print("write: ", filename)

laba_file = path_to_project + '/' + name_solution + '\\'

for filename in glob.glob(laba_file + "*.cpp"):
    print("filename:", filename)
    if (filename not in glob.glob(laba_file + "laba_*.cpp")) and (filename not in glob.glob(laba_file + "ConsoleAlgorithms.cpp")):
        addCodeTitle("\nCode: " + filename[filename.find('\\') : ] + '\n', p)
        addCode(open(filename, "r").read(), p)
        print("_write: ", filename)

for filename in glob.glob(laba_file + "*.h"):
    print("filename:", filename)
    if  filename not in glob.glob(laba_file + "laba_*.h"):
        addCodeTitle("\nCode: " + filename[filename.find('\\') : ] + '\n', p)
        addCode(open(filename, "r").read(), p)
        print("_write:", filename)


doc.save('test_.docx')